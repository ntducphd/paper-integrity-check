# -*- coding: utf-8 -*-
"""
paper-integrity-check -- a local pre-check for a researcher's own manuscript portfolio.

This is NOT a Turnitin/iThenticate replacement -- those tools work because they pay
publishers for access to a proprietary database of billions of web pages and journal
articles, which nothing here has access to. What IS fully checkable from files already
on your machine, and is specifically valuable if you run several parallel manuscripts on
related topics at once:

  1. self-plagiarism / text recycling -- does a manuscript share long passages with ANY
     OTHER manuscript or already-published paper in your own portfolio (every file under
     --root, plus any --extra-dir you point at your own published-paper archives)?
  2. citation-paraphrase -- does a manuscript share long passages with the full text of a
     source it cites, when that source's PDF is sitting locally in a project's
     literature/ folder (i.e., quoting a cited paper too closely instead of paraphrasing)?
  3. ai-tells -- known AI-writing lexical/structural patterns (stock phrases, overused
     constructions, em-dash/colon clustering), formalized into one reusable script.

This script by itself does NOT check against the open web or any publisher's database of
OTHER authors' work you have not cited and do not have a local copy of. The `candidates`
command below scans the whole manuscript for that gap (every sufficiently distinctive
sentence, not a sample -- see WEB_CHECK.md for the full Stage 3 procedure, which needs a
Claude Code session to actually search them via WebSearch and compile a verdict per
sentence; `stage3-report` then renders that compiled verdict list as a visual HTML report).
Treat all of this as a pre-check that catches the risks unique to running many parallel
manuscripts on similar topics, run it alongside (not instead of) your institution's real
Turnitin/iThenticate access.

Usage:
    python integrity_check.py build-cache --root <your-papers-folder> [--rebuild]
                                           [--extra-dir PATH ...] [--exclude TOKEN ...]
    python integrity_check.py check <path-to-manuscript.docx-or-.md> --root <your-papers-folder>
                                     [--rebuild-cache] [--extra-dir PATH ...] [--exclude TOKEN ...]
                                     [--max-doc-freq N] [--min-words N]
    python integrity_check.py candidates <path-to-manuscript.docx-or-.md> [--top N]
                                          [--min-words N] [--max-words N]
    python integrity_check.py stage3-report <compiled-results.json>

Output: a Markdown report written to reports/<stem>_integrity_report_<date>.md
(plus .json/.html siblings for `check`; a candidate-sentence .json for `candidates`;
an .html dashboard for `stage3-report`)
"""
import argparse
import hashlib
import json
import re
import sys
from collections import defaultdict
from datetime import date
from pathlib import Path

import fitz  # pymupdf
from docx import Document

HERE = Path(__file__).resolve().parent
CACHE_FILE = HERE / "cache" / "corpus_cache.json"
REPORTS_DIR = HERE / "reports"

SHINGLE_SIZE = 8            # words per shingle
MIN_PASSAGE_WORDS = 20       # minimum contiguous overlap to flag as a real passage, not noise
MAX_DOC_FREQ_FOR_FLAG = 3    # a shingle shared by more than this many OTHER docs is boilerplate, not flagged

DEFAULT_EXCLUDE_DIR_TOKENS = {
    "archive", "backup", ".venv", "venv", "node_modules", "__pycache__",
    ".git", "cache", "reports",
}
EXTS = {".docx", ".pdf", ".md"}

WORD_RE = re.compile(r"\w+", re.UNICODE)


# --------------------------------------------------------------------------- extraction

def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def extract_pdf_text(path: Path) -> str:
    parts = []
    with fitz.open(str(path)) as d:
        for page in d:
            parts.append(page.get_text())
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suf = path.suffix.lower()
    try:
        if suf == ".docx":
            return extract_docx_text(path)
        if suf == ".pdf":
            return extract_pdf_text(path)
        if suf in (".md", ".txt"):
            return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"  [skip, extraction failed] {path.name}: {e}", file=sys.stderr)
    return ""


def tokenize(text: str):
    return [w.lower() for w in WORD_RE.findall(text)]


def paragraphs_of(text: str):
    return [p for p in re.split(r"\n+", text) if p.strip()]


# --------------------------------------------------------------------------- shingling

def stable_hash(s: str) -> int:
    return int(hashlib.blake2b(s.encode("utf-8"), digest_size=8).hexdigest(), 16)


def make_shingles(tokens, n=SHINGLE_SIZE):
    shingles = defaultdict(list)
    for i in range(len(tokens) - n + 1):
        gram = " ".join(tokens[i:i + n])
        shingles[stable_hash(gram)].append(i)
    return shingles


def merge_passages(pos_pairs):
    """Group (target_pos, other_pos) pairs into contiguous diagonal runs."""
    by_diag = defaultdict(list)
    for t, o in pos_pairs:
        by_diag[o - t].append((t, o))
    passages = []
    for pts in by_diag.values():
        pts = sorted(set(pts))
        start_t, start_o = pts[0]
        prev_t, prev_o = pts[0]
        for t, o in pts[1:]:
            if t == prev_t + 1:
                prev_t, prev_o = t, o
                continue
            passages.append((start_t, prev_t, start_o, prev_o))
            start_t, start_o = t, o
            prev_t, prev_o = t, o
        passages.append((start_t, prev_t, start_o, prev_o))
    return passages


# --------------------------------------------------------------------------- corpus

def classify_path(p: Path) -> str:
    parts_lower = [x.lower() for x in p.parts]
    if "literature" in parts_lower:
        return "cited_source"
    return "own_manuscript"


def norm_key(p: Path) -> str:
    return str(p.resolve()).replace("\\", "/").lower()


def is_excluded(p: Path, exclude_tokens) -> bool:
    if p.name.startswith("~$"):  # Word lock files
        return True
    try:
        if p.resolve().is_relative_to(HERE):  # never scan the tool's own folder
            return True
    except (OSError, ValueError):
        pass
    parts_lower = [x.lower() for x in p.parts]
    return any(tok in part for part in parts_lower for tok in exclude_tokens)


def iter_corpus_files(root: Path, extra_dirs, exclude_tokens):
    roots = [(root, None)]
    for d in extra_dirs:
        d = Path(d)
        if d.exists():
            roots.append((d, "own_published"))
        else:
            print(f"  [--extra-dir not found, skipped] {d}", file=sys.stderr)
    for base, forced_category in roots:
        for p in base.rglob("*"):
            if not p.is_file() or p.suffix.lower() not in EXTS:
                continue
            if is_excluded(p, exclude_tokens):
                continue
            yield p, (forced_category or classify_path(p))


def dedupe_docs(docs: dict) -> int:
    """Mark byte-identical extracted text (e.g. the same manuscript copied into a
    submission_package/ folder) as a duplicate of one canonical entry, so the report
    doesn't repeat the same match 3 times for 3 copies of the same file. Marked (not
    deleted) so a re-run's mtime check still recognizes the file and skips re-extracting
    it -- deleting the entry made every run pay the extraction cost again for every
    duplicate, every time."""
    seen = {}
    n_marked = 0
    for key in sorted(docs.keys()):
        d = docs[key]
        if "shingles" not in d:
            continue  # already marked as a duplicate in a prior run
        h = hashlib.sha256(" ".join(d["tokens"]).encode("utf-8")).hexdigest()
        if h in seen:
            d["duplicate_of"] = seen[h]
            d.pop("tokens", None)
            d.pop("shingles", None)
            n_marked += 1
        else:
            seen[h] = key
    return n_marked


def build_cache(root: Path, extra_dirs=(), exclude_tokens=None, rebuild=False) -> dict:
    if exclude_tokens is None:
        exclude_tokens = DEFAULT_EXCLUDE_DIR_TOKENS
    cache = {"docs": {}}
    if CACHE_FILE.exists() and not rebuild:
        cache = json.loads(CACHE_FILE.read_text(encoding="utf-8"))
    docs = cache["docs"]
    seen_keys = set()
    n_new = 0
    n_seen = 0
    if extra_dirs:
        print(f"Including {len(extra_dirs)} extra dir(s) -- if any are on a cloud-synced drive "
              f"(Google Drive, OneDrive, etc.), the first scan can be very slow (each file may "
              f"trigger a network fetch).")
    for p, category in iter_corpus_files(root, extra_dirs, exclude_tokens):
        key = norm_key(p)
        seen_keys.add(key)
        n_seen += 1
        if n_seen % 100 == 0:
            print(f"  ...scanned {n_seen} files ({n_new} extracted so far): {p.name}")
        try:
            mtime = p.stat().st_mtime
        except OSError:
            continue
        cached = docs.get(key)
        if cached and abs(cached.get("mtime", -1) - mtime) < 1e-6:
            continue
        text = extract_text(p)
        tokens = tokenize(text)
        if len(tokens) < SHINGLE_SIZE:
            continue
        shingles = make_shingles(tokens)
        docs[key] = {
            "display_path": str(p),
            "mtime": mtime,
            "category": category,
            "tokens": tokens,
            "shingles": {str(h): pos for h, pos in shingles.items()},
        }
        n_new += 1

    if not extra_dirs:
        # don't evict cached extra-dir entries just because this run didn't pass --extra-dir
        seen_keys |= {k for k, d in docs.items() if d.get("category") == "own_published"}

    for key in list(docs.keys()):
        if key not in seen_keys:
            del docs[key]

    n_dupe = dedupe_docs(docs)

    CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
    CACHE_FILE.write_text(json.dumps(cache, ensure_ascii=False), encoding="utf-8")
    print(f"Corpus cache: {len(docs)} documents "
          f"({n_seen} files scanned, {n_new} newly extracted, {n_dupe} exact-duplicate copies collapsed).")
    return cache


# --------------------------------------------------------------------------- checking

BIBLIO_MARKERS = ("doi org", "https", "www.")
BARE_DOI_RE = re.compile(r"\bdoi\s+10\s+\d{4}\b")


def looks_like_bibliography(snippet: str) -> bool:
    """Heuristic: a matched passage that is mostly a citation string (author/year/journal/DOI)
    rather than prose. Verified against this tool's own first real report: every single
    citation-paraphrase match it found was this kind of shared reference-list entry, not
    copied original prose -- exactly the same pattern this portfolio's manual self-plagiarism
    checks have found by hand every time. Flagging it separately, rather than silently
    dropping it, keeps it auditable.

    Extended 2026-07-25 after a Stage-2 agent review (see AGENT_REVIEW.md) found this heuristic
    was missing citation strings that sit in an INFORMAL context -- a markdown literature-notes
    file, or a PDF-cache extraction -- rather than a formatted Word reference-list entry. Those
    snippets often carry only ONE year (not 2+) and no DOI/URL within the matched window, but
    still read as "surname initial, surname initial, ..." -- an author-initial list, tokenized
    down to a run of single-character tokens ("t", "v", "g", "p", ...). That pattern is itself a
    strong, distinct signal, checked in addition to (not instead of) the original two checks.

    Extended again 2026-07-25 (same day) after a real cross-manuscript run found a shared
    reference entry formatted as "DOI: 10.1007/..." (no "doi.org" domain, so BIBLIO_MARKERS
    missed it) between two companion papers -- a bare "doi 10 nnnn" token sequence is itself
    unambiguous, checked regardless of the other signals."""
    low = snippet.lower()
    if any(m in low for m in BIBLIO_MARKERS):
        return True
    if BARE_DOI_RE.search(low):
        return True
    if len(re.findall(r"\b(19|20)\d{2}\b", snippet)) >= 2:
        return True
    words = snippet.split()
    if len(words) < 6:
        return False
    n_single_char = sum(1 for w in words if len(w) == 1)
    return n_single_char >= 3 and (n_single_char / len(words)) >= 0.15


SENTENCE_END_RE = re.compile(r"(?<=[.!?])\s+(?=[A-Z(])")
_ABBREV_TAIL_RE = re.compile(
    r"\b(?:et al|e\.g|i\.e|fig|eq|vs|dr|prof|approx|cf|ref|no|vol|pp)\.$", re.IGNORECASE)


def split_sentences(paragraph: str):
    """Rough sentence splitter -- good enough for picking Stage 3 web-search candidates, not
    meant to be a precise NLP sentence tokenizer. Fragments ending in a common academic
    abbreviation ("et al.", "Fig.", "e.g.", ...) are merged back into the next fragment so
    those aren't mistaken for sentence boundaries."""
    raw = SENTENCE_END_RE.split(paragraph.strip())
    sentences, buf = [], ""
    for frag in raw:
        buf = f"{buf} {frag}".strip() if buf else frag
        if _ABBREV_TAIL_RE.search(buf):
            continue
        sentences.append(buf)
        buf = ""
    if buf:
        sentences.append(buf)
    return [s for s in sentences if s]


def select_candidates(text: str, top_n=None, min_words: int = 12, max_words: int = 40):
    """Pick the manuscript's "googleable" sentences for Stage 3 (see WEB_CHECK.md): long
    enough to be distinctive rather than generic, not a heading/table fragment, not already a
    citation string (reuses looks_like_bibliography -- no point web-searching a reference-list
    entry). By default (top_n=None) this returns EVERY qualifying sentence -- a full-manuscript
    scan, not a sample -- sorted longest-first so a capped run (--top N) still gets the most
    distinctive ones first. A 12-word floor and 40-word ceiling are the only exclusions: shorter
    sentences match too many unrelated documents to be a useful query, and 40+-word ones are
    usually a run-on or a list rather than a single searchable claim -- not a coverage gap in
    the plagiarism-detection sense, since neither extreme is where copied prose is diagnosable.

    This is the deterministic, fully-offline half of Stage 3. It does NOT search anything itself
    -- turning this list into actual web/database hits requires either a Claude Code session's
    WebSearch tool (see WEB_CHECK.md) or a paid search API this tool does not assume you have."""
    seen = set()
    candidates = []
    for para in paragraphs_of(text):
        # Check bibliography-likeness at the WHOLE-PARAGRAPH level first, before splitting into
        # sentences. A reference-list entry ("Surname A., Surname B. (2023). Title of the paper.
        # Journal. 12(3): 45.") trips the author-initial-list / year heuristics when read as one
        # unit, but sentence-splitting on the period after "(2023)." separates the author/year
        # signal from the title fragment -- and a bare title alone often carries NONE of those
        # signals, so it silently passed the per-sentence check and got offered up as a "distinctive
        # manuscript sentence" to search, when it's actually just the reference list. Caught via a
        # real Stage 3 run: a WebSearch agent flagged a paper's own title (reused verbatim in ITS
        # OWN reference-list entry, as expected) as a suspicious "HIT" because, out of context, it
        # looked like unquoted prose reusing someone else's title.
        if looks_like_bibliography(" ".join(tokenize(para))):
            continue
        for sent in split_sentences(para):
            n_words = len(sent.split())
            if not (min_words <= n_words <= max_words):
                continue
            # looks_like_bibliography() expects tokenized (punctuation-stripped) text -- that's
            # how Stage 1 always calls it (on tokenize()'d shingle snippets). Passing raw sentence
            # text with punctuation attached defeats its author-initial-list check: "M.," and "M"
            # are different strings, so a real author list like "Baseggio M., Murray M., ..." was
            # slipping through here even though the same words tokenized would trip the detector.
            if looks_like_bibliography(" ".join(tokenize(sent))):
                continue
            key = sent.lower()
            if key in seen:
                continue
            seen.add(key)
            candidates.append({"sentence": sent, "n_words": n_words})
    candidates.sort(key=lambda c: -c["n_words"])
    return candidates[:top_n]


def same_document_family(target_path: Path, other_display_path: str) -> bool:
    """A manuscript's own .docx/.pdf/.md exports (and its own README) living in the SAME
    folder are the same document, not a different one to flag as self-plagiarism -- exclude
    them. Genuine cross-manuscript reuse (a different project, or even a different manuscript
    subfolder of the same project) is NOT excluded by this check."""
    other = Path(other_display_path)
    return other.parent.resolve() == target_path.parent.resolve()


def find_overlaps(target_path, target_tokens, target_shingles, docs, target_key, max_doc_freq, min_words):
    doc_freq = defaultdict(set)
    for key, d in docs.items():
        if key == target_key or "shingles" not in d:  # duplicate-marked entries carry no shingles
            continue
        for h in d["shingles"].keys():
            doc_freq[int(h)].add(key)

    matches = defaultdict(list)  # other_key -> [(target_pos, other_pos), ...]
    for h, positions in target_shingles.items():
        others = doc_freq.get(h)
        if not others or len(others) > max_doc_freq:
            continue
        for other_key in others:
            if same_document_family(target_path, docs[other_key]["display_path"]):
                continue
            other_positions = docs[other_key]["shingles"].get(str(h), [])
            for tp in positions:
                for op in other_positions:
                    matches[other_key].append((tp, op))

    results = []
    for other_key, pos_pairs in matches.items():
        for (t_start, t_end, o_start, o_end) in merge_passages(pos_pairs):
            n_words = t_end - t_start + SHINGLE_SIZE
            if n_words < min_words:
                continue
            snippet = " ".join(target_tokens[t_start:t_end + SHINGLE_SIZE])
            results.append({
                "other_key": other_key,
                "display_path": docs[other_key]["display_path"],
                "category": docs[other_key]["category"],
                "n_words": n_words,
                "snippet": snippet,
                "is_biblio": looks_like_bibliography(snippet),
                "t_start": t_start,
                "t_end": t_end + SHINGLE_SIZE - 1,  # inclusive token index, for highlighting
            })
    results.sort(key=lambda r: -r["n_words"])
    return results


# --------------------------------------------------------------------------- AI tells

AI_PHRASES = [
    "delve", "leverage", "tapestry", "underscore", "holistic", "seamlessly", "seamless",
    "paves the way", "sheds light on", "garner", "bolster", "myriad", "plethora",
    "intricate", "nuanced", "multifaceted", "testament to", "unwavering", "cutting-edge",
    "game-changer", "game changer", "unlock the potential", "harness the power",
    "in today's world", "at the forefront", "navigate the complexities",
    "it is important to note", "it is worth noting", "crucial role", "vital role",
    "pivotal role", "in conclusion", "not only", "furthermore", "moreover",
    "boasts", "underpin", "foster a", "fostering",
]


def scan_ai_tells(text: str, tokens_count: int):
    findings = {}
    lower = text.lower()
    phrase_hits = {}
    for phrase in AI_PHRASES:
        n = lower.count(phrase)
        if n:
            phrase_hits[phrase] = n
    findings["phrases"] = phrase_hits

    findings["comma_not_construction"] = len(re.findall(r",\s*not\s+\w+", text, flags=re.IGNORECASE))

    paras = paragraphs_of(text)
    # A markdown table separator/rule row (e.g. "|---|---|---|") is pure table
    # syntax, not prose -- counting its "--" runs as an em-dash-style AI tell
    # is a false positive that scales with how many tables a manuscript has.
    md_table_rule_re = re.compile(r"^\s*\|?[\s\-:|]+\|?\s*$")
    em_dash_clusters = [
        i for i, p in enumerate(paras)
        if not md_table_rule_re.match(p) and (p.count("\u2014") + p.count("--")) >= 2
    ]
    colon_clusters = [i for i, p in enumerate(paras) if p.count(":") >= 2]
    findings["em_dash_cluster_paragraphs"] = len(em_dash_clusters)
    findings["colon_cluster_paragraphs"] = len(colon_clusters)
    findings["total_paragraphs"] = len(paras)
    findings["total_words"] = tokens_count
    return findings


# --------------------------------------------------------------------------- reporting

def render_matches_section(lines, matches, group_label):
    if not matches:
        lines.append(f"No matching passages found against any {group_label}.")
        lines.append("")
        return

    prose = [m for m in matches if not m["is_biblio"]]
    biblio = [m for m in matches if m["is_biblio"]]

    if prose:
        lines.append(f"**{len(prose)} passage(s) of real prose overlap** — review these individually:")
        lines.append("")
        by_doc = defaultdict(list)
        for m in prose:
            by_doc[m["display_path"]].append(m)
        for doc_path, ms in sorted(by_doc.items(), key=lambda kv: -sum(x["n_words"] for x in kv[1])):
            lines.append(f"#### vs. `{doc_path}`")
            for m in sorted(ms, key=lambda x: -x["n_words"]):
                lines.append(f"- **{m['n_words']} words**: “{m['snippet']}”")
            lines.append("")
    else:
        lines.append("No prose overlap found (only shared reference-list entries below, if any).")
        lines.append("")

    if biblio:
        by_doc = defaultdict(list)
        for m in biblio:
            by_doc[m["display_path"]].append(m)
        lines.append(f"<details><summary>{len(biblio)} shared reference-list / citation-string "
                      f"match(es) — both documents cite the same sources, this is expected and is "
                      f"NOT a prose-reuse concern (collapsed by default)</summary>")
        lines.append("")
        for doc_path, ms in sorted(by_doc.items(), key=lambda kv: -sum(x["n_words"] for x in kv[1])):
            lines.append(f"- `{doc_path}`: {len(ms)} shared entries, longest {max(x['n_words'] for x in ms)} words")
        lines.append("")
        lines.append("</details>")
        lines.append("")


def write_report(target_path: Path, self_matches, cited_matches, ai_findings, max_doc_freq, min_words):
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stem = target_path.stem
    out_path = REPORTS_DIR / f"{stem}_integrity_report_{date.today().isoformat()}.md"

    lines = []
    lines.append(f"# Academic-integrity check: {target_path.name}")
    lines.append("")
    lines.append(f"Run: {date.today().isoformat()} | Target: `{target_path}`")
    lines.append("")
    lines.append("> **Scope note**: this checks your own portfolio (all manuscripts under "
                  "`D:/2026_Paper` + your published-paper archives) and locally-available cited-source "
                  "PDFs (`literature/` folders). It does **not** query the open web or any publisher "
                  "database, so it cannot catch overlap with a paper you did not cite and do not have a "
                  "local copy of — that gap is exactly what a real Turnitin/iThenticate submission covers. "
                  "Use this as a pre-check for the risks specific to running many parallel manuscripts on "
                  "related topics, not a substitute for institutional plagiarism screening.")
    lines.append("")

    lines.append(f"## 1. Self-plagiarism / text reuse across your own portfolio "
                 f"(passages >={min_words} words, shared with <={max_doc_freq} other documents; "
                 f"the target's own same-folder .docx/.pdf/.md/README siblings are excluded automatically)")
    lines.append("")
    render_matches_section(lines, self_matches, "other manuscript or published paper in the portfolio")

    lines.append(f"## 2. Over-close paraphrase of a cited source's own text "
                 f"(passages >={min_words} words, matched against local `literature/` PDFs)")
    lines.append("")
    render_matches_section(lines, cited_matches, "locally-available cited-source text")

    lines.append("## 3. AI-writing-tell scan")
    lines.append("")
    lines.append(f"- Total words: {ai_findings['total_words']}, paragraphs: {ai_findings['total_paragraphs']}")
    lines.append(f"- `, not X` construction count: {ai_findings['comma_not_construction']}")
    lines.append(f"- Paragraphs with 2+ em-dashes: {ai_findings['em_dash_cluster_paragraphs']}")
    lines.append(f"- Paragraphs with 2+ colons: {ai_findings['colon_cluster_paragraphs']}")
    if ai_findings["phrases"]:
        lines.append("- Stock AI phrases found:")
        for phrase, n in sorted(ai_findings["phrases"].items(), key=lambda kv: -kv[1]):
            lines.append(f"  - \"{phrase}\": {n}")
    else:
        lines.append("- No stock AI phrases from the checklist found.")
    lines.append("")

    lines.append("## Summary")
    lines.append("")
    n_self = len(self_matches)
    n_cited = len(cited_matches)
    n_ai = ai_findings['comma_not_construction'] + ai_findings['em_dash_cluster_paragraphs'] + \
        ai_findings['colon_cluster_paragraphs'] + sum(ai_findings["phrases"].values())
    lines.append(f"{n_self} self-reuse passage(s), {n_cited} close-paraphrase passage(s), "
                 f"{n_ai} AI-tell hit(s) (raw count across all categories above). "
                 f"Review each flagged passage individually — a shared passage may be a legitimate, "
                 f"already-disclosed self-citation of your own prior method description rather than a "
                 f"problem; this tool flags candidates, it does not adjudicate them.")

    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Report written: {out_path}")
    return out_path


def write_json_report(target_path: Path, self_matches, cited_matches, ai_findings):
    """Machine-readable sibling to the Markdown report -- the input format Stage 2 (agent-based
    semantic review, see AGENT_REVIEW.md) reads. Only the 'real prose' matches are worth sending
    to an agent; bibliography-entry matches are already explained and not a semantic question."""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stem = target_path.stem
    out_path = REPORTS_DIR / f"{stem}_integrity_report_{date.today().isoformat()}.json"

    def slim(matches):
        prose = [m for m in matches if not m["is_biblio"]]
        biblio = [m for m in matches if m["is_biblio"]]
        return {
            "prose_matches": [
                {"display_path": m["display_path"], "category": m["category"],
                 "n_words": m["n_words"], "snippet": m["snippet"]}
                for m in prose
            ],
            "biblio_match_count": len(biblio),
        }

    payload = {
        "target": str(target_path),
        "run_date": date.today().isoformat(),
        "self_plagiarism": slim(self_matches),
        "citation_paraphrase": slim(cited_matches),
        "ai_tells": ai_findings,
    }
    out_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"JSON report written: {out_path}")
    return out_path


# --------------------------------------------------------------------------- HTML report

def compute_overlap_index(n_tokens: int, self_matches, cited_matches) -> float:
    """Percentage of the target document's own words that fall inside at least one
    non-bibliography flagged passage. Deliberately named 'Local Overlap Index', not
    'similarity score' -- unlike Turnitin's number, this is scoped ONLY to --root and
    --extra-dir, so a low score here says nothing about overlap with the wider published
    literature. See the report's own methodology note before quoting this number out of
    context."""
    covered = set()
    for m in self_matches + cited_matches:
        if m["is_biblio"]:
            continue
        covered.update(range(m["t_start"], m["t_end"] + 1))
    return 100.0 * len(covered) / n_tokens if n_tokens else 0.0


def _source_breakdown(matches, n_tokens):
    """Rank sources by total flagged (non-biblio) words, for the dashboard table --
    the local equivalent of Turnitin's per-source similarity breakdown."""
    by_doc = defaultdict(lambda: {"words": 0, "passages": 0, "category": None})
    for m in matches:
        if m["is_biblio"]:
            continue
        d = by_doc[m["display_path"]]
        d["words"] += m["n_words"]
        d["passages"] += 1
        d["category"] = m["category"]
    rows = [
        {"path": path, "words": v["words"], "passages": v["passages"],
         "category": v["category"], "pct": 100.0 * v["words"] / n_tokens if n_tokens else 0.0}
        for path, v in by_doc.items()
    ]
    rows.sort(key=lambda r: -r["words"])
    return rows


def _build_highlight_spans(target_tokens, self_matches, cited_matches):
    """One category label per token index ('self' / 'cited' / None), self takes priority
    on the rare token covered by both, so the reconstructed document can be rendered with
    one <mark> per contiguous run instead of one per matched shingle."""
    labels = [None] * len(target_tokens)
    for m in cited_matches:
        if m["is_biblio"]:
            continue
        for i in range(m["t_start"], m["t_end"] + 1):
            labels[i] = ("cited", m["display_path"])
    for m in self_matches:  # self overwrites cited on overlap -- self-reuse is the higher-priority flag
        if m["is_biblio"]:
            continue
        for i in range(m["t_start"], m["t_end"] + 1):
            labels[i] = ("self", m["display_path"])
    return labels


def render_highlighted_body(text: str, token_spans, labels) -> str:
    """Render the ORIGINAL document text -- real casing, punctuation, paragraph breaks --
    as HTML, wrapping the character span of each flagged token run in a <mark>. token_spans[i]
    is the (start, end) character offset of token i within `text` (from WORD_RE.finditer, which
    -- for a plain \\w+ pattern with no groups -- yields matches in the same order/count as the
    WORD_RE.findall used by tokenize(), so token indices line up exactly). Rendering the real
    text rather than the lowercased token stream is what makes this readable like an actual
    plagiarism-checker report instead of a word-soup dump."""
    import html as _html
    paras = []
    start = 0
    for line in text.split("\n"):
        end = start + len(line)
        if line.strip():
            paras.append((start, end))
        start = end + 1  # +1 for the '\n' separator consumed by split()

    out_paras = []
    tok_i = 0
    n_tok = len(token_spans)
    for p_start, p_end in paras:
        while tok_i < n_tok and token_spans[tok_i][0] < p_start:
            tok_i += 1
        j = tok_i
        while j < n_tok and token_spans[j][0] < p_end:
            j += 1
        para_tok_idx = list(range(tok_i, j))
        tok_i = j

        pieces = []
        cursor = p_start
        k = 0
        m = len(para_tok_idx)
        while k < m:
            idx = para_tok_idx[k]
            lab = labels[idx]
            k2 = k
            while k2 < m and labels[para_tok_idx[k2]] == lab:
                k2 += 1
            run_start_tok, run_end_tok = para_tok_idx[k], para_tok_idx[k2 - 1]
            span_start, span_end = token_spans[run_start_tok][0], token_spans[run_end_tok][1]
            if span_start > cursor:
                pieces.append(_html.escape(text[cursor:span_start]))
            if lab is None:
                pieces.append(_html.escape(text[span_start:span_end]))
            else:
                cat, src = lab
                title = _html.escape(f"Matches: {src}")
                pieces.append(f'<mark class="hl-{cat}" title="{title}">'
                               f'{_html.escape(text[span_start:span_end])}</mark>')
            cursor = span_end
            k = k2
        if cursor < p_end:
            pieces.append(_html.escape(text[cursor:p_end]))
        out_paras.append(f"<p>{''.join(pieces)}</p>")
    return "\n".join(out_paras)


HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Integrity check: {title}</title>
<style>
  :root {{
    --bg: #faf9f6; --panel: #ffffff; --ink: #1f1b16; --ink-soft: #5b554c;
    --line: #e4ddd0; --accent: #8a5a2b;
    --self: #f4c98a; --self-ink: #6b4415;
    --cited: #a9c9e0; --cited-ink: #1c3a54;
  }}
  @media (prefers-color-scheme: dark) {{
    :root {{
      --bg: #17140f; --panel: #221e17; --ink: #f1ede4; --ink-soft: #b8b0a1;
      --line: #3a3327; --accent: #e0a860;
      --self: #6b4a1e; --self-ink: #ffdca3;
      --cited: #2c4a63; --cited-ink: #bfe0ff;
    }}
  }}
  :root[data-theme="dark"] {{
    --bg: #17140f; --panel: #221e17; --ink: #f1ede4; --ink-soft: #b8b0a1;
    --line: #3a3327; --accent: #e0a860;
    --self: #6b4a1e; --self-ink: #ffdca3;
    --cited: #2c4a63; --cited-ink: #bfe0ff;
  }}
  :root[data-theme="light"] {{
    --bg: #faf9f6; --panel: #ffffff; --ink: #1f1b16; --ink-soft: #5b554c;
    --line: #e4ddd0; --accent: #8a5a2b;
    --self: #f4c98a; --self-ink: #6b4415;
    --cited: #a9c9e0; --cited-ink: #1c3a54;
  }}
  * {{ box-sizing: border-box; }}
  body {{
    margin: 0; background: var(--bg); color: var(--ink);
    font: 16px/1.55 -apple-system, "Segoe UI", Helvetica, Arial, sans-serif;
  }}
  .wrap {{ max-width: 920px; margin: 0 auto; padding: 2.5rem 1.5rem 5rem; }}
  h1 {{ font-size: 1.5rem; margin: 0 0 .25rem; text-wrap: balance; }}
  .meta {{ color: var(--ink-soft); font-size: .9rem; margin-bottom: 1.75rem; }}
  .scope {{
    background: var(--panel); border: 1px solid var(--line); border-left: 4px solid var(--accent);
    border-radius: 6px; padding: 1rem 1.25rem; font-size: .92rem; color: var(--ink-soft);
    margin-bottom: 2rem;
  }}
  .scope strong {{ color: var(--ink); }}
  .dashboard {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 1px;
    background: var(--line); border: 1px solid var(--line); border-radius: 8px; overflow: hidden;
    margin-bottom: 2rem; }}
  .card {{ background: var(--panel); padding: 1.1rem 1.25rem; }}
  .card .num {{ font-size: 2rem; font-weight: 700; font-variant-numeric: tabular-nums; line-height: 1; }}
  .card .lbl {{ color: var(--ink-soft); font-size: .78rem; text-transform: uppercase; letter-spacing: .04em;
    margin-top: .4rem; }}
  h2 {{ font-size: 1.05rem; margin: 2.25rem 0 .75rem; padding-bottom: .4rem; border-bottom: 1px solid var(--line); }}
  table {{ width: 100%; border-collapse: collapse; font-size: .88rem; }}
  th, td {{ text-align: left; padding: .5rem .6rem; border-bottom: 1px solid var(--line); }}
  th {{ color: var(--ink-soft); font-weight: 600; font-size: .78rem; text-transform: uppercase; letter-spacing: .03em; }}
  td.num, th.num {{ text-align: right; font-variant-numeric: tabular-nums; white-space: nowrap; }}
  .src-path {{ font-family: ui-monospace, Consolas, monospace; font-size: .8rem; word-break: break-all; }}
  .tag {{ display: inline-block; font-size: .72rem; padding: .1rem .45rem; border-radius: 999px; font-weight: 600; }}
  .tag-self {{ background: var(--self); color: var(--self-ink); }}
  .tag-cited {{ background: var(--cited); color: var(--cited-ink); }}
  .legend {{ display: flex; gap: 1.25rem; font-size: .82rem; color: var(--ink-soft); margin-bottom: .75rem; flex-wrap: wrap; }}
  .legend span {{ display: inline-flex; align-items: center; gap: .35rem; }}
  .swatch {{ width: .8rem; height: .8rem; border-radius: 3px; display: inline-block; }}
  .doc-body {{
    background: var(--panel); border: 1px solid var(--line); border-radius: 8px;
    padding: 1.5rem 1.75rem; font-family: Georgia, "Times New Roman", serif; font-size: .96rem;
    line-height: 1.75; max-width: 68ch; overflow-wrap: break-word;
  }}
  mark.hl-self {{ background: var(--self); color: var(--self-ink); padding: .05rem .1rem; border-radius: 3px; cursor: help; }}
  mark.hl-cited {{ background: var(--cited); color: var(--cited-ink); padding: .05rem .1rem; border-radius: 3px; cursor: help; }}
  .ai-tells {{ background: var(--panel); border: 1px solid var(--line); border-radius: 8px; padding: 1.25rem 1.5rem; }}
  .ai-tells ul {{ margin: .5rem 0 0; padding-left: 1.2rem; }}
  .empty {{ color: var(--ink-soft); font-style: italic; }}
  footer {{ margin-top: 3rem; color: var(--ink-soft); font-size: .8rem; border-top: 1px solid var(--line); padding-top: 1rem; }}
</style>
</head>
<body>
<div class="wrap">
  <h1>Academic-integrity check: {title}</h1>
  <div class="meta">Run {run_date} &middot; local corpus scan, {n_docs} documents indexed</div>

  <div class="scope"><strong>Scope.</strong> This checks your own portfolio and locally-available
    cited-source PDFs only &mdash; it does not query the open web or any publisher database, so it
    cannot catch overlap with a paper you never cited and do not have a local copy of. That gap is
    exactly what a real Turnitin/iThenticate submission covers; run this alongside it, not instead
    of it. The Local Overlap Index below is <em>not</em> a Turnitin similarity score &mdash; it is
    scoped strictly to the documents this tool was pointed at.</div>

  <div class="dashboard">
    <div class="card"><div class="num">{overlap_index:.1f}%</div><div class="lbl">Local overlap index</div></div>
    <div class="card"><div class="num">{n_self_src}</div><div class="lbl">Self-reuse sources</div></div>
    <div class="card"><div class="num">{n_cited_src}</div><div class="lbl">Close-paraphrase sources</div></div>
    <div class="card"><div class="num">{n_words}</div><div class="lbl">Document words</div></div>
    <div class="card"><div class="num">{n_ai_hits}</div><div class="lbl">AI-tell hits</div></div>
  </div>

  <h2>Matched sources</h2>
  {source_table}

  <h2>Document, with matches highlighted</h2>
  <div class="legend">
    <span><span class="swatch" style="background:var(--self)"></span> self-plagiarism candidate (another manuscript in your portfolio)</span>
    <span><span class="swatch" style="background:var(--cited)"></span> over-close paraphrase of a cited source</span>
  </div>
  <div class="doc-body">{highlighted_body}</div>

  <h2>AI-writing-tell scan</h2>
  <div class="ai-tells">
    <table>
      <tr><td>&#8220;, not X&#8221; construction</td><td class="num">{comma_not}</td></tr>
      <tr><td>Paragraphs with 2+ em-dashes</td><td class="num">{em_dash_paras}</td></tr>
      <tr><td>Paragraphs with 2+ colons</td><td class="num">{colon_paras}</td></tr>
    </table>
    {ai_phrase_list}
  </div>

  <footer>Generated by <a href="https://github.com/ntducphd/paper-integrity-check">paper-integrity-check</a>
    &middot; Stage 1 (lexical scan) only &mdash; see AGENT_REVIEW.md for the optional agent-based semantic
    review layer &middot; not a substitute for institutional Turnitin/iThenticate screening.</footer>
</div>
</body>
</html>
"""


def render_html_report(target_path: Path, target_tokens, self_matches, cited_matches, ai_findings,
                        n_docs_indexed: int) -> Path:
    import html as _html
    n_tokens = len(target_tokens)
    overlap_index = compute_overlap_index(n_tokens, self_matches, cited_matches)

    src_rows = _source_breakdown(self_matches, n_tokens) + _source_breakdown(cited_matches, n_tokens)
    src_rows.sort(key=lambda r: -r["words"])
    if src_rows:
        body_rows = []
        for r in src_rows:
            tag = '<span class="tag tag-self">self</span>' if r["category"] in ("own_manuscript", "own_published") \
                else '<span class="tag tag-cited">cited</span>'
            body_rows.append(
                f'<tr><td>{tag}</td><td class="src-path">{_html.escape(r["path"])}</td>'
                f'<td class="num">{r["words"]}</td><td class="num">{r["passages"]}</td>'
                f'<td class="num">{r["pct"]:.1f}%</td></tr>'
            )
        source_table = ('<table><tr><th></th><th>Source</th><th class="num">Words</th>'
                         '<th class="num">Passages</th><th class="num">% of doc</th></tr>'
                         + "".join(body_rows) + "</table>")
    else:
        source_table = '<p class="empty">No non-bibliography matches against any source.</p>'

    labels = _build_highlight_spans(target_tokens, self_matches, cited_matches)
    original_text = extract_text(target_path)
    token_spans = [(m.start(), m.end()) for m in WORD_RE.finditer(original_text)]
    highlighted_body = render_highlighted_body(original_text, token_spans, labels)

    if ai_findings["phrases"]:
        items = "".join(f"<li>&#8220;{_html.escape(p)}&#8221; &times; {n}</li>"
                         for p, n in sorted(ai_findings["phrases"].items(), key=lambda kv: -kv[1]))
        ai_phrase_list = f"<p>Stock AI phrases found:</p><ul>{items}</ul>"
    else:
        ai_phrase_list = '<p class="empty">No stock AI phrases from the checklist found.</p>'

    n_self_src = len({r["path"] for r in _source_breakdown(self_matches, n_tokens)})
    n_cited_src = len({r["path"] for r in _source_breakdown(cited_matches, n_tokens)})
    n_ai_hits = (ai_findings["comma_not_construction"] + ai_findings["em_dash_cluster_paragraphs"]
                 + ai_findings["colon_cluster_paragraphs"] + sum(ai_findings["phrases"].values()))

    html_out = HTML_TEMPLATE.format(
        title=_html.escape(target_path.name), run_date=date.today().isoformat(),
        n_docs=n_docs_indexed, overlap_index=overlap_index, n_self_src=n_self_src,
        n_cited_src=n_cited_src, n_words=n_tokens, n_ai_hits=n_ai_hits,
        source_table=source_table, highlighted_body=highlighted_body,
        comma_not=ai_findings["comma_not_construction"],
        em_dash_paras=ai_findings["em_dash_cluster_paragraphs"],
        colon_paras=ai_findings["colon_cluster_paragraphs"],
        ai_phrase_list=ai_phrase_list,
    )

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    out_path = REPORTS_DIR / f"{target_path.stem}_integrity_report_{date.today().isoformat()}.html"
    out_path.write_text(html_out, encoding="utf-8")
    print(f"HTML report written: {out_path}")
    return out_path


STAGE3_HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Stage 3 open-web check: {title}</title>
<style>
  :root {{
    --bg: #faf9f6; --panel: #ffffff; --ink: #1f1b16; --ink-soft: #5b554c;
    --line: #e4ddd0; --accent: #8a5a2b;
    --clean: #bfe0c8; --clean-ink: #185c34;
    --hit: #f0b8ae; --hit-ink: #7a2418;
  }}
  @media (prefers-color-scheme: dark) {{
    :root {{
      --bg: #17140f; --panel: #221e17; --ink: #f1ede4; --ink-soft: #b8b0a1;
      --line: #3a3327; --accent: #e0a860;
      --clean: #1f4a30; --clean-ink: #9fe0b8;
      --hit: #5a251c; --hit-ink: #ffc0b0;
    }}
  }}
  :root[data-theme="dark"] {{
    --bg: #17140f; --panel: #221e17; --ink: #f1ede4; --ink-soft: #b8b0a1;
    --line: #3a3327; --accent: #e0a860;
    --clean: #1f4a30; --clean-ink: #9fe0b8;
    --hit: #5a251c; --hit-ink: #ffc0b0;
  }}
  :root[data-theme="light"] {{
    --bg: #faf9f6; --panel: #ffffff; --ink: #1f1b16; --ink-soft: #5b554c;
    --line: #e4ddd0; --accent: #8a5a2b;
    --clean: #bfe0c8; --clean-ink: #185c34;
    --hit: #f0b8ae; --hit-ink: #7a2418;
  }}
  * {{ box-sizing: border-box; }}
  body {{
    margin: 0; background: var(--bg); color: var(--ink);
    font: 16px/1.55 -apple-system, "Segoe UI", Helvetica, Arial, sans-serif;
  }}
  .wrap {{ max-width: 920px; margin: 0 auto; padding: 2.5rem 1.5rem 5rem; }}
  h1 {{ font-size: 1.5rem; margin: 0 0 .25rem; text-wrap: balance; }}
  .meta {{ color: var(--ink-soft); font-size: .9rem; margin-bottom: 1.75rem; }}
  .scope {{
    background: var(--panel); border: 1px solid var(--line); border-left: 4px solid var(--accent);
    border-radius: 6px; padding: 1rem 1.25rem; font-size: .92rem; color: var(--ink-soft);
    margin-bottom: 2rem;
  }}
  .scope strong {{ color: var(--ink); }}
  .dashboard {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 1px;
    background: var(--line); border: 1px solid var(--line); border-radius: 8px; overflow: hidden;
    margin-bottom: 2rem; }}
  .card {{ background: var(--panel); padding: 1.1rem 1.25rem; }}
  .card .num {{ font-size: 2rem; font-weight: 700; font-variant-numeric: tabular-nums; line-height: 1; }}
  .card .lbl {{ color: var(--ink-soft); font-size: .78rem; text-transform: uppercase; letter-spacing: .04em;
    margin-top: .4rem; }}
  h2 {{ font-size: 1.05rem; margin: 2.25rem 0 .75rem; padding-bottom: .4rem; border-bottom: 1px solid var(--line); }}
  table {{ width: 100%; border-collapse: collapse; font-size: .88rem; }}
  th, td {{ text-align: left; padding: .5rem .6rem; border-bottom: 1px solid var(--line); vertical-align: top; }}
  th {{ color: var(--ink-soft); font-weight: 600; font-size: .78rem; text-transform: uppercase; letter-spacing: .03em; }}
  td.num, th.num {{ text-align: right; font-variant-numeric: tabular-nums; white-space: nowrap; }}
  .tag {{ display: inline-block; font-size: .72rem; padding: .1rem .5rem; border-radius: 999px; font-weight: 600; white-space: nowrap; }}
  .tag-clean {{ background: var(--clean); color: var(--clean-ink); }}
  .tag-hit {{ background: var(--hit); color: var(--hit-ink); }}
  .hit-note {{ font-size: .85rem; color: var(--ink-soft); margin-top: .3rem; }}
  .sentence-list {{ background: var(--panel); border: 1px solid var(--line); border-radius: 8px;
    padding: .25rem 1.5rem; }}
  .empty {{ color: var(--ink-soft); font-style: italic; }}
  footer {{ margin-top: 3rem; color: var(--ink-soft); font-size: .8rem; border-top: 1px solid var(--line); padding-top: 1rem; }}
</style>
</head>
<body>
<div class="wrap">
  <h1>Stage 3: open-web cross-check &mdash; {title}</h1>
  <div class="meta">Run {run_date} &middot; {n_scanned} sentence(s) searched via WebSearch (exact-phrase query per sentence)</div>

  <div class="scope"><strong>What this is.</strong> Stage 1 only compares files already on this machine.
    Stage 3 searches the manuscript's distinctive sentences against the open web &mdash; the same thing
    as copying a sentence into Google, done for every sentence long enough to be worth searching, not a
    sample. It does not cover Scopus or Web of Science (both require a paid institutional API key not
    available here) and a clean result means nothing indexed by the search backend used currently
    matches &mdash; not an absolute guarantee. See WEB_CHECK.md for the full scope note.</div>

  <div class="dashboard">
    <div class="card"><div class="num">{n_scanned}</div><div class="lbl">Sentences scanned</div></div>
    <div class="card"><div class="num">{n_clean}</div><div class="lbl">Clean</div></div>
    <div class="card"><div class="num">{n_hit}</div><div class="lbl">Flagged for review</div></div>
    <div class="card"><div class="num">{pct_clean:.0f}%</div><div class="lbl">Clean rate</div></div>
  </div>

  {pipeline_note}

  <h2>Flagged for review</h2>
  {hits_section}

  <h2>All scanned sentences</h2>
  <div class="sentence-list">
    <table>
      <tr><th style="width:3rem">#</th><th></th><th>Sentence</th><th class="num">Words</th></tr>
      {all_rows}
    </table>
  </div>

  <footer>Generated by <a href="https://github.com/ntducphd/paper-integrity-check">paper-integrity-check</a>
    &middot; Stage 3 (open-web cross-check via a Claude Code session's WebSearch tool) &mdash; see
    WEB_CHECK.md for the full procedure &middot; not a substitute for institutional
    Turnitin/iThenticate screening.</footer>
</div>
</body>
</html>
"""


def render_stage3_html_report(target_path: Path, results: list, pipeline_note: str = "",
                               out_path: Path = None) -> Path:
    """Render Stage 3's (WEB_CHECK.md) compiled WebSearch results as a self-contained HTML
    report, visually consistent with Stage 1's render_html_report (same design language) so
    the two feel like one tool, not two bolted-together scripts. `results` is a list of
    {"sentence": str, "n_words": int, "verdict": "clean"|"hit", "note": str} dicts -- Stage 3
    itself is agent-driven (a Claude Code session runs WebSearch and judges each result), so
    this function only renders an already-compiled verdict list, it does not search anything."""
    import html as _html
    n_scanned = len(results)
    n_hit = sum(1 for r in results if r["verdict"] == "hit")
    n_clean = n_scanned - n_hit
    pct_clean = 100.0 * n_clean / n_scanned if n_scanned else 0.0

    hits = [r for r in results if r["verdict"] == "hit"]
    if hits:
        hit_rows = "".join(
            f'<tr><td><span class="tag tag-hit">hit</span></td>'
            f'<td>{_html.escape(r["sentence"])}<div class="hit-note">{_html.escape(r.get("note", ""))}</div></td></tr>'
            for r in hits
        )
        hits_section = f"<table>{hit_rows}</table>"
    else:
        hits_section = '<p class="empty">No sentence returned a verbatim/near-verbatim match against an unrelated source.</p>'

    row_parts = []
    for i, r in enumerate(results, 1):
        tag = '<span class="tag tag-hit">hit</span>' if r["verdict"] == "hit" else '<span class="tag tag-clean">clean</span>'
        row_parts.append(
            f'<tr><td class="num">{i}</td><td>{tag}</td>'
            f'<td>{_html.escape(r["sentence"])}</td><td class="num">{r["n_words"]}</td></tr>'
        )
    all_rows = "".join(row_parts)

    pipeline_note_html = (
        f'<div class="scope"><strong>Pipeline transparency note.</strong> {pipeline_note}</div>'
        if pipeline_note else ""
    )

    html_out = STAGE3_HTML_TEMPLATE.format(
        title=_html.escape(target_path.name), run_date=date.today().isoformat(),
        n_scanned=n_scanned, n_clean=n_clean, n_hit=n_hit, pct_clean=pct_clean,
        pipeline_note=pipeline_note_html, hits_section=hits_section, all_rows=all_rows,
    )

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    out_path = out_path or REPORTS_DIR / f"{target_path.stem}_stage3_web_check_{date.today().isoformat()}.html"
    out_path.write_text(html_out, encoding="utf-8")
    print(f"Stage 3 HTML report written: {out_path}")
    return out_path


# --------------------------------------------------------------------------- CLI

def run_check(target_path: Path, cache: dict, max_doc_freq: int, min_words: int):
    docs = cache["docs"]
    target_key = norm_key(target_path)
    text = extract_text(target_path)
    tokens = tokenize(text)
    shingles = make_shingles(tokens)

    all_matches = find_overlaps(target_path, tokens, shingles, docs, target_key, max_doc_freq, min_words)
    self_matches = [m for m in all_matches if m["category"] in ("own_manuscript", "own_published")]
    cited_matches = [m for m in all_matches if m["category"] == "cited_source"]

    ai_findings = scan_ai_tells(text, len(tokens))
    return self_matches, cited_matches, ai_findings, tokens


def add_corpus_args(p):
    p.add_argument("--root", type=Path, default=Path.cwd(),
                    help="Folder to scan for your manuscripts (default: current directory)")
    p.add_argument("--extra-dir", action="append", default=[], dest="extra_dirs",
                    help="Additional folder to scan, tagged as your own published work "
                         "(e.g. a separate published-papers archive). Repeatable. Off by "
                         "default -- pass explicitly, since a cloud-synced drive can be slow "
                         "to scan the first time.")
    p.add_argument("--exclude", action="append", default=[], dest="extra_excludes",
                    help="Extra path-component token to exclude from scanning (e.g. a "
                         f"personal website repo folder name). Repeatable. Defaults already "
                         f"excluded: {sorted(DEFAULT_EXCLUDE_DIR_TOKENS)}")


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_build = sub.add_parser("build-cache", help="Build or refresh the corpus cache")
    p_build.add_argument("--rebuild", action="store_true", help="Force full re-extraction of every file")
    add_corpus_args(p_build)

    p_check = sub.add_parser("check", help="Run integrity checks on one manuscript")
    p_check.add_argument("target", help="Path to the manuscript (.docx or .md)")
    p_check.add_argument("--rebuild-cache", action="store_true")
    add_corpus_args(p_check)
    p_check.add_argument("--max-doc-freq", type=int, default=MAX_DOC_FREQ_FOR_FLAG,
                          help="A passage shared with more than this many other documents is treated as "
                               "boilerplate and not flagged (default: %(default)s)")
    p_check.add_argument("--min-words", type=int, default=MIN_PASSAGE_WORDS,
                          help="Minimum contiguous matched passage length in words (default: %(default)s)")

    p_cand = sub.add_parser(
        "candidates",
        help="Pick a manuscript's distinctive sentences for Stage 3 web/database search "
             "(see WEB_CHECK.md) -- offline, no corpus needed. Outputs ALL qualifying "
             "sentences by default (a full-manuscript scan, not a sample); use --top to cap it.")
    p_cand.add_argument("target", help="Path to the manuscript (.docx or .md)")
    p_cand.add_argument("--top", type=int, default=None,
                         help="Cap the number of candidate sentences output. Default: no cap "
                              "-- every qualifying sentence in the manuscript, for a full scan.")
    p_cand.add_argument("--min-words", type=int, default=12,
                         help="Minimum sentence length in words to be considered (default: %(default)s). "
                              "Sentences shorter than this are too generic to search meaningfully -- "
                              "excluded from the scan for that reason, not for speed.")
    p_cand.add_argument("--max-words", type=int, default=40,
                         help="Maximum sentence length in words to be considered (default: %(default)s)")

    p_stage3 = sub.add_parser(
        "stage3-report",
        help="Render a compiled Stage 3 WebSearch results file (see WEB_CHECK.md) as a "
             "visual HTML report, matching Stage 1's report design")
    p_stage3.add_argument(
        "results", type=Path,
        help='Path to a JSON file: {"target": "<manuscript path>", "pipeline_note": "<optional>", '
             '"results": [{"sentence": str, "n_words": int, "verdict": "clean"|"hit", "note": str}, ...]}')

    args = parser.parse_args()

    if args.cmd == "build-cache":
        exclude_tokens = DEFAULT_EXCLUDE_DIR_TOKENS | set(args.extra_excludes)
        build_cache(args.root, extra_dirs=args.extra_dirs, exclude_tokens=exclude_tokens,
                    rebuild=args.rebuild)
    elif args.cmd == "check":
        exclude_tokens = DEFAULT_EXCLUDE_DIR_TOKENS | set(args.extra_excludes)
        target = Path(args.target).resolve()
        if not target.exists():
            sys.exit(f"Not found: {target}")
        cache = build_cache(args.root, extra_dirs=args.extra_dirs, exclude_tokens=exclude_tokens,
                             rebuild=args.rebuild_cache)
        self_matches, cited_matches, ai_findings, target_tokens = run_check(
            target, cache, args.max_doc_freq, args.min_words)
        write_report(target, self_matches, cited_matches, ai_findings, args.max_doc_freq, args.min_words)
        write_json_report(target, self_matches, cited_matches, ai_findings)
        render_html_report(target, target_tokens, self_matches, cited_matches, ai_findings, len(cache["docs"]))
    elif args.cmd == "candidates":
        target = Path(args.target).resolve()
        if not target.exists():
            sys.exit(f"Not found: {target}")
        text = extract_text(target)
        cands = select_candidates(text, top_n=args.top, min_words=args.min_words, max_words=args.max_words)
        REPORTS_DIR.mkdir(parents=True, exist_ok=True)
        out_path = REPORTS_DIR / f"{target.stem}_candidates_{date.today().isoformat()}.json"
        out_path.write_text(json.dumps({"target": str(target), "candidates": cands},
                                        ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"{len(cands)} candidate sentence(s) written: {out_path}\n")
        for i, c in enumerate(cands, 1):
            print(f"{i}. ({c['n_words']}w) {c['sentence']}")
    elif args.cmd == "stage3-report":
        if not args.results.exists():
            sys.exit(f"Not found: {args.results}")
        payload = json.loads(args.results.read_text(encoding="utf-8"))
        target = Path(payload["target"])
        render_stage3_html_report(target, payload["results"], payload.get("pipeline_note", ""))


if __name__ == "__main__":
    main()
